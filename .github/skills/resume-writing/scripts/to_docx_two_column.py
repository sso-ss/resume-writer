#!/usr/bin/env python3
"""Convert a Markdown resume to a two-column .docx — Layout B (sidebar-left).

Layout B: Name anchored inside left sidebar
────────────────────────────────────────────
┌────────────────┬─────────────────────────────────────┐
│  JENNIFER      │                                     │
│  LAUREN        │  SUMMARY                            │
│  ─────────     │  3 lines full width in main col     │
│  CONTACT       │                                     │
│  Portfolio     │  EXPERIENCE                         │
│  LinkedIn      │  Role | Company | Dates             │
│  Email         │  • bullet                           │
│  Location      │  • bullet                           │
│  ─────────     │                                     │
│  SKILLS        │  Role | Company | Dates             │
│  Design: ...   │  • bullet                           │
│  Research: ... │  • bullet                           │
│  ─────────     │                                     │
│  EDUCATION     │  KEY PROJECTS                       │
│  Degree, Uni   │  ■ Project + link                   │
│  ─────────     │  ■ Project + link                   │
│  TOOLS         │                                     │
│  Figma, ...    │                                     │
└────────────────┴─────────────────────────────────────┘

Name lives inside sidebar — anchored, not floating. Everything static
in the left, everything narrative on the right.

Usage:
    python to_docx_two_column.py <input.md> [output.docx]
    python to_docx_two_column.py <input.md> --layout b   (explicit)

Requires: python-docx (pip install python-docx)
"""

from __future__ import annotations

import re
import sys
from pathlib import Path
from typing import Dict, List, Optional, Tuple

try:
    from docx import Document
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("Error: python-docx is required. Install it with: pip install python-docx")
    sys.exit(1)

# -- Design tokens --
ACCENT = RGBColor(0x2B, 0x4C, 0x5E)
SIDEBAR_BG = "F0EFEB"
TEXT_DARK = RGBColor(0x1A, 0x1A, 0x1A)
TEXT_BODY = RGBColor(0x33, 0x33, 0x33)
TEXT_MUTED = RGBColor(0x66, 0x66, 0x66)
FONT = "Calibri"

SIDEBAR_W = Inches(2.4)
MAIN_W = Inches(4.8)


# ---------------------------------------------------------------------------
# Shared helpers
# ---------------------------------------------------------------------------

def parse_markdown(md_text: str) -> Tuple[str, str, Dict[str, List[str]]]:
    lines = md_text.splitlines()
    name = ""
    contact = ""
    sections: Dict[str, List[str]] = {}
    current = ""
    for line in lines:
        s = line.rstrip()
        if s.startswith("# ") and not s.startswith("## "):
            name = s[2:].strip()
        elif name and not contact and s and not s.startswith("#"):
            contact = s.strip()
        elif s.startswith("## "):
            current = s[3:].strip().lower()
            sections[current] = []
        elif current:
            sections[current].append(s)
    return name, contact, sections


def parse_contact_items(contact_line: str) -> List[Tuple[str, str]]:
    items: List[Tuple[str, str]] = []
    for part in contact_line.split("|"):
        part = part.strip()
        if not part:
            continue
        if ":" in part and not part.startswith("http"):
            label, _, value = part.partition(":")
            items.append((label.strip(), value.strip()))
        elif "@" in part:
            items.append(("Email", part))
        elif "linkedin.com" in part.lower():
            items.append(("LinkedIn", part))
        elif "http" in part.lower():
            items.append(("Portfolio", part))
        else:
            items.append(("Location", part))
    return items


def _set_cell_shading(cell, hex_color: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _set_cell_margins(cell, top=0, start=0, bottom=0, end=0) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    margins = OxmlElement("w:tcMar")
    for edge, val in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        m = OxmlElement(f"w:{edge}")
        m.set(qn("w:w"), str(val))
        m.set(qn("w:type"), "dxa")
        margins.append(m)
    tcPr.append(margins)


def _remove_table_borders(table) -> None:
    tblPr = table._tbl.tblPr
    if tblPr is None:
        return
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        elem = OxmlElement(f"w:{edge}")
        elem.set(qn("w:val"), "nil")
        borders.append(elem)
    tblPr.append(borders)


def _fmt(paragraph, text: str, size=Pt(9.5), color=TEXT_BODY, bold=False):
    """Render text with **bold** and [link](url) flattening."""
    for part in re.split(r"(\*\*.*?\*\*)", text):
        if part.startswith("**") and part.endswith("**"):
            r = paragraph.add_run(part[2:-2])
            r.bold = True
            r.font.size = size
            r.font.color.rgb = color
            r.font.name = FONT
        else:
            for frag in re.split(r"\[([^\]]+)\]\([^\)]+\)", part):
                if frag:
                    r = paragraph.add_run(frag)
                    r.bold = bold
                    r.font.size = size
                    r.font.color.rgb = color
                    r.font.name = FONT


def _heading(cell, text: str, accent=ACCENT) -> None:
    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text.upper())
    r.bold = True
    r.font.size = Pt(8.5)
    r.font.color.rgb = accent
    r.font.name = FONT
    rPr = r._r.get_or_add_rPr()
    sp = OxmlElement("w:spacing")
    sp.set(qn("w:val"), "30")
    rPr.append(sp)


def _sidebar_lines(cell, lines: List[str]) -> None:
    for raw in lines:
        line = raw.strip()
        if not line:
            continue
        bullet = re.match(r"^[\-*]\s+(.*)", line)
        if bullet:
            p = cell.add_paragraph()
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.left_indent = Inches(0.05)
            _fmt(p, bullet.group(1).strip(), size=Pt(8.5))
        elif line.startswith("### "):
            p = cell.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(1)
            r = p.add_run(line[4:].strip())
            r.bold = True
            r.font.size = Pt(9)
            r.font.color.rgb = TEXT_DARK
            r.font.name = FONT
        else:
            p = cell.add_paragraph()
            p.paragraph_format.space_after = Pt(1)
            _fmt(p, line, size=Pt(8.5))


def _main_lines(cell, lines: List[str]) -> None:
    for raw in lines:
        line = raw.strip()
        if not line:
            continue
        if line.startswith("### "):
            p = cell.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(2)
            _fmt(p, line[4:].strip(), size=Pt(10), color=TEXT_DARK, bold=True)
        elif re.match(r"^[\-*]\s+", line):
            content = re.match(r"^[\-*]\s+(.*)", line).group(1).strip()
            p = cell.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.left_indent = Inches(0.15)
            p.clear()
            _fmt(p, content, size=Pt(9))
        else:
            p = cell.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            _fmt(p, line, size=Pt(9.5))


# ---------------------------------------------------------------------------
# Layout B builder
# ---------------------------------------------------------------------------

def build_layout_b(md_path: str, docx_path: Optional[str] = None) -> str:
    md_file = Path(md_path)
    if not md_file.exists():
        print(f"Error: {md_path} not found")
        sys.exit(1)

    if docx_path is None:
        base = md_file.stem.replace("_Resume", "")
        docx_path = str(md_file.with_name(f"{base}_ProductDesigner_Resume.docx"))

    name, contact_line, sections = parse_markdown(md_file.read_text(encoding="utf-8"))
    contact_items = parse_contact_items(contact_line) if contact_line else []

    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.4)
    sec.bottom_margin = Inches(0.4)
    sec.left_margin = Inches(0.35)
    sec.right_margin = Inches(0.35)

    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = Pt(10)
    style.font.color.rgb = TEXT_BODY
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.space_before = Pt(0)

    # Full-width table, one row
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    sidebar, main = table.rows[0].cells
    sidebar.width = SIDEBAR_W
    main.width = MAIN_W

    _set_cell_shading(sidebar, SIDEBAR_BG)
    _set_cell_margins(sidebar, top=140, start=140, bottom=140, end=100)
    _set_cell_margins(main, top=140, start=180, bottom=140, end=80)

    # ── Sidebar: Name inside ──
    # Split name into first/last for stacked display
    parts = (name or "Full Name").split(None, 1)
    first_name = parts[0] if parts else ""
    last_name = parts[1] if len(parts) > 1 else ""

    p = sidebar.paragraphs[0]  # use the default empty paragraph
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(first_name)
    r.bold = True
    r.font.size = Pt(20)
    r.font.color.rgb = TEXT_DARK
    r.font.name = FONT

    if last_name:
        p = sidebar.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run(last_name)
        r.bold = True
        r.font.size = Pt(20)
        r.font.color.rgb = TEXT_DARK
        r.font.name = FONT

    # Thin separator
    p = sidebar.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("─" * 18)
    r.font.size = Pt(6)
    r.font.color.rgb = TEXT_MUTED
    r.font.name = FONT

    # Contact
    if contact_items:
        _heading(sidebar, "Contact")
        for label, value in contact_items:
            p = sidebar.add_paragraph()
            p.paragraph_format.space_after = Pt(1)
            r = p.add_run(f"{label}\n")
            r.font.size = Pt(7)
            r.font.color.rgb = TEXT_MUTED
            r.font.name = FONT
            r.bold = True
            r = p.add_run(value)
            r.font.size = Pt(8.5)
            r.font.color.rgb = TEXT_DARK
            r.font.name = FONT

    # Skills
    for key in ("skills & tools", "skills"):
        if key in sections:
            _heading(sidebar, "Skills")
            _sidebar_lines(sidebar, sections[key])
            break

    # Separator
    p = sidebar.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("─" * 18)
    r.font.size = Pt(6)
    r.font.color.rgb = TEXT_MUTED

    # Education
    if "education" in sections:
        _heading(sidebar, "Education")
        _sidebar_lines(sidebar, sections["education"])

    # ── Main column ──
    # Clear the default empty paragraph
    main.paragraphs[0].clear()

    if "summary" in sections:
        _heading(main, "Summary")
        _main_lines(main, sections["summary"])

    if "experience" in sections:
        _heading(main, "Experience")
        _main_lines(main, sections["experience"])

    for key in ("key projects", "projects"):
        if key in sections:
            _heading(main, "Key Projects")
            _main_lines(main, sections[key])
            break

    if "recognition" in sections:
        _heading(main, "Recognition")
        _main_lines(main, sections["recognition"])

    if "volunteer" in sections:
        _heading(main, "Volunteer")
        _main_lines(main, sections["volunteer"])

    _remove_table_borders(table)
    doc.save(docx_path)
    return docx_path


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python to_docx_two_column.py <input.md> [output.docx]")
        sys.exit(1)

    input_path = sys.argv[1]
    output_path = sys.argv[2] if len(sys.argv) > 2 else None
    result = build_layout_b(input_path, output_path)
    print(f"Layout B (sidebar-left) resume saved to: {result}")

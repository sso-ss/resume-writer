#!/usr/bin/env python3
"""Convert a Markdown resume to a professionally formatted .docx file.

Usage:
    python to_docx.py <input.md> [output.docx]

If output is omitted, replaces .md with .docx in the filename.
Requires: python-docx (pip install python-docx)
"""

from __future__ import annotations

import sys
import re
from pathlib import Path
from typing import Optional

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.section import WD_ORIENT
except ImportError:
    print("Error: python-docx is required. Install it with: pip install python-docx")
    sys.exit(1)


def create_resume_doc(md_path: str, docx_path: Optional[str] = None) -> str:
    """Parse a Markdown resume and produce a formatted .docx file."""
    md_file = Path(md_path)
    if not md_file.exists():
        print(f"Error: {md_path} not found")
        sys.exit(1)

    if docx_path is None:
        docx_path = str(md_file.with_suffix(".docx"))

    content = md_file.read_text(encoding="utf-8")
    lines = content.split("\n")

    doc = Document()

    # -- Page setup --
    section = doc.sections[0]
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)

    # -- Default font --
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(10)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    style.paragraph_format.space_after = Pt(2)
    style.paragraph_format.space_before = Pt(0)

    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        # --- H1: Name (centered, large, bold) ---
        if line.startswith("# "):
            name = line[2:].strip()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(name)
            run.bold = True
            run.font.size = Pt(18)
            run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
            i += 1
            continue

        # --- Contact line (line right after name, no heading marker) ---
        # Detect contact-like lines: contain | separators or multiple links
        if (
            i > 0
            and not line.startswith("#")
            and "|" in line
            and any(
                kw in line.lower()
                for kw in ["@", "linkedin", "portfolio", "http", "gmail", "email"]
            )
        ):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(6)
            _add_formatted_run(p, line.strip(), size=Pt(9), color=RGBColor(0x55, 0x55, 0x55))
            i += 1
            continue

        # --- H2: Section headings ---
        if line.startswith("## "):
            heading_text = line[3:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(3)
            run = p.add_run(heading_text.upper())
            run.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
            # Add a bottom border (thin line)
            _add_bottom_border(p)
            i += 1
            continue

        # --- H3: Subsection (role title / company line) ---
        if line.startswith("### "):
            sub_text = line[4:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(2)
            _add_formatted_run(p, sub_text, size=Pt(10), bold=True, color=RGBColor(0x1A, 0x1A, 0x1A))
            i += 1
            continue

        # --- Bullet points ---
        bullet_match = re.match(r"^[\-\*]\s+(.*)", line)
        if bullet_match:
            bullet_text = bullet_match.group(1).strip()
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.left_indent = Inches(0.25)
            # Clear default run and add formatted text
            p.clear()
            _add_formatted_run(p, bullet_text, size=Pt(9.5))
            i += 1
            continue

        # --- Regular text (summary, project descriptions, etc.) ---
        if line.strip():
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            _add_formatted_run(p, line.strip(), size=Pt(10))
            i += 1
            continue

        # Blank line
        i += 1

    doc.save(docx_path)
    return docx_path


def _add_formatted_run(paragraph, text: str, size=None, bold=False, color=None):
    """Add a run with inline bold (**text**) support."""
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            # Handle [text](url) links — render as plain text (ATS-safe)
            link_parts = re.split(r"\[([^\]]+)\]\([^\)]+\)", part)
            for lp in link_parts:
                if lp:
                    run = paragraph.add_run(lp)
                    run.bold = bold
        if size:
            for run in paragraph.runs:
                run.font.size = size
        if color:
            for run in paragraph.runs:
                run.font.color.rgb = color


def _add_bottom_border(paragraph):
    """Add a thin bottom border to a paragraph for section headings."""
    from docx.oxml.ns import qn
    from lxml import etree

    pPr = paragraph._p.get_or_add_pPr()
    pBdr = etree.SubElement(pPr, qn("w:pBdr"))
    bottom = etree.SubElement(pBdr, qn("w:bottom"))
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "AAAAAA")


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python to_docx.py <input.md> [output.docx]")
        sys.exit(1)

    input_path = sys.argv[1]
    output_path = sys.argv[2] if len(sys.argv) > 2 else None
    result = create_resume_doc(input_path, output_path)
    print(f"Resume saved to: {result}")

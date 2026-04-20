"""Shared helpers for resume .docx conversion scripts.

Used by to_docx_two_column.py (Layout B) and to_docx_right_sidebar.py (Layout C).
"""

from __future__ import annotations

import re
from typing import Dict, List, Tuple

from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

# -- Design tokens --
ACCENT = RGBColor(0x2B, 0x4C, 0x5E)
SIDEBAR_BG = "F0EFEB"
TEXT_DARK = RGBColor(0x1A, 0x1A, 0x1A)
TEXT_BODY = RGBColor(0x33, 0x33, 0x33)
TEXT_MUTED = RGBColor(0x66, 0x66, 0x66)
FONT = "Calibri"

SIDEBAR_W = Inches(2.4)
MAIN_W = Inches(4.8)


def parse_markdown(md_text: str) -> Tuple[str, str, Dict[str, List[str]]]:
    """Parse a Markdown resume into name, contact line, and sections dict."""
    lines = md_text.splitlines()
    name = ""
    contact = ""
    sections: Dict[str, List[str]] = {}
    current = ""
    found_name = False
    for line in lines:
        s = line.rstrip()
        if s.startswith("# ") and not s.startswith("## "):
            name = s[2:].strip()
            found_name = True
        elif found_name and not contact and not s.startswith("#"):
            # Skip blank lines between name and contact
            if not s:
                continue
            contact = s.strip()
        elif s.startswith("## "):
            current = s[3:].strip().lower()
            sections[current] = []
        elif current:
            sections[current].append(s)
    return name, contact, sections


def parse_contact_items(contact_line: str) -> List[Tuple[str, str]]:
    """Split a pipe-delimited contact line into (label, value) tuples."""
    items: List[Tuple[str, str]] = []
    for part in contact_line.split("|"):
        part = part.strip()
        if not part:
            continue
        if ":" in part and not part.startswith("http"):
            label, _, value = part.partition(":")
            items.append((label.strip(), value.strip()))
        elif "@" in part:
            items.append(("Email", part))
        elif "linkedin.com" in part.lower():
            items.append(("LinkedIn", part))
        elif "http" in part.lower():
            items.append(("Portfolio", part))
        else:
            items.append(("Location", part))
    return items


def set_cell_shading(cell, hex_color: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_margins(cell, top=0, start=0, bottom=0, end=0) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    margins = OxmlElement("w:tcMar")
    for edge, val in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        m = OxmlElement(f"w:{edge}")
        m.set(qn("w:w"), str(val))
        m.set(qn("w:type"), "dxa")
        margins.append(m)
    tcPr.append(margins)


def remove_table_borders(table) -> None:
    tblPr = table._tbl.tblPr
    if tblPr is None:
        return
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        elem = OxmlElement(f"w:{edge}")
        elem.set(qn("w:val"), "nil")
        borders.append(elem)
    tblPr.append(borders)


def add_bottom_rule(paragraph) -> None:
    """Add a thin bottom border to a paragraph."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "3")
    bottom.set(qn("w:color"), "AAAAAA")
    pBdr.append(bottom)
    pPr.append(pBdr)


def fmt(paragraph, text: str, size=Pt(9.5), color=TEXT_BODY, bold=False):
    """Render text with **bold** and [link](url) flattening."""
    for part in re.split(r"(\*\*.*?\*\*)", text):
        if part.startswith("**") and part.endswith("**"):
            r = paragraph.add_run(part[2:-2])
            r.bold = True
            r.font.size = size
            r.font.color.rgb = color
            r.font.name = FONT
        else:
            for frag in re.split(r"\[([^\]]+)\]\([^\)]+\)", part):
                if frag:
                    r = paragraph.add_run(frag)
                    r.bold = bold
                    r.font.size = size
                    r.font.color.rgb = color
                    r.font.name = FONT


def heading(cell, text: str, accent=ACCENT) -> None:
    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text.upper())
    r.bold = True
    r.font.size = Pt(8.5)
    r.font.color.rgb = accent
    r.font.name = FONT
    rPr = r._r.get_or_add_rPr()
    sp = OxmlElement("w:spacing")
    sp.set(qn("w:val"), "30")
    rPr.append(sp)


def sidebar_lines(cell, lines: List[str]) -> None:
    for raw in lines:
        line = raw.strip()
        if not line:
            continue
        bullet = re.match(r"^[\-*]\s+(.*)", line)
        if bullet:
            p = cell.add_paragraph()
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.left_indent = Inches(0.05)
            fmt(p, bullet.group(1).strip(), size=Pt(8.5))
        elif line.startswith("### "):
            p = cell.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(1)
            r = p.add_run(line[4:].strip())
            r.bold = True
            r.font.size = Pt(9)
            r.font.color.rgb = TEXT_DARK
            r.font.name = FONT
        else:
            p = cell.add_paragraph()
            p.paragraph_format.space_after = Pt(1)
            fmt(p, line, size=Pt(8.5))


def main_lines(cell, lines: List[str]) -> None:
    for raw in lines:
        line = raw.strip()
        if not line:
            continue
        if line.startswith("### "):
            p = cell.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(2)
            fmt(p, line[4:].strip(), size=Pt(10), color=TEXT_DARK, bold=True)
        elif re.match(r"^[\-*]\s+", line):
            content = re.match(r"^[\-*]\s+(.*)", line).group(1).strip()
            p = cell.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.left_indent = Inches(0.15)
            p.clear()
            fmt(p, content, size=Pt(9))
        else:
            p = cell.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            fmt(p, line, size=Pt(9.5))


def render_sidebar_skills(sidebar, sections: Dict[str, List[str]]) -> None:
    """Render Skills and Tools sections into sidebar cell."""
    for key in ("skills & tools", "skills"):
        if key in sections:
            heading(sidebar, "Skills")
            sidebar_lines(sidebar, sections[key])
            break

    # Tools (standalone section, if not combined with Skills)
    if "tools" in sections and "skills & tools" not in sections:
        heading(sidebar, "Tools")
        sidebar_lines(sidebar, sections["tools"])


def render_main_content(main, sections: Dict[str, List[str]]) -> None:
    """Render main content sections (Summary, Experience, Projects, etc.)."""
    main.paragraphs[0].clear()

    if "summary" in sections:
        heading(main, "Summary")
        main_lines(main, sections["summary"])

    if "experience" in sections:
        heading(main, "Experience")
        main_lines(main, sections["experience"])

    for key in ("key projects", "projects"):
        if key in sections:
            heading(main, "Key Projects")
            main_lines(main, sections[key])
            break

    if "recognition" in sections:
        heading(main, "Recognition")
        main_lines(main, sections["recognition"])

    if "volunteer" in sections:
        heading(main, "Volunteer")
        main_lines(main, sections["volunteer"])
